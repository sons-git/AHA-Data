import csv
import io
import base64
from docx import Document
from PyPDF2 import PdfReader
from typing import List, Tuple, Optional
from app.schemas.conversations import FileData, ProcessedMessage
from app.utils.image_processing import convert_to_dspy_image

def extract_text(file_data: FileData) -> Optional[str]:
    """
    Extract text content from various file types based on FileData.file.

    Args:
        file_data (FileData): Metadata and content of the file.

    Returns:
        Optional[str]: Extracted text content, or None if extraction fails.
    """
    try:
        # Decode bytes from base64 if needed
        if isinstance(file_data.file, str):
            content_bytes = base64.b64decode(file_data.file.split(",", 1)[-1])
        else:
            content_bytes = file_data.file

        # Handle plain text and markdown files
        if file_data.type in {"text/plain", "text/markdown"}:
            return content_bytes.decode("utf-8", errors="ignore")

        # Handle CSV files
        elif file_data.type == "text/csv":
            decoded = content_bytes.decode("utf-8", errors="ignore")
            reader = csv.reader(io.StringIO(decoded))
            rows = [" | ".join(row) for row in reader]
            return "\n".join(rows) if rows else None

        # Handle PDF files
        elif file_data.type == "application/pdf":
            reader = PdfReader(io.BytesIO(content_bytes))
            if reader.is_encrypted:
                raise ValueError(f"PDF file '{file_data.name}' is encrypted and cannot be processed.")
            extracted = []
            for page_num, page in enumerate(reader.pages):
                try:
                    text = page.extract_text()
                    if text:
                        extracted.append(text)
                except Exception as e:
                    print(f"Failed to extract text from page {page_num} in {file_data.name}: {e}")
            return "\n".join(extracted) if extracted else None

        # Handle DOCX files
        elif file_data.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(io.BytesIO(content_bytes))
            paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
            tables = []
            for table in doc.tables:
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if cells:
                        tables.append(" | ".join(cells))
            all_text = paragraphs + tables
            return "\n".join(all_text) if all_text else None

    except Exception as e:
        print(f"Failed to extract text from {file_data.name}: {e}")
        return None

def classify_file(files: List[FileData]) -> Tuple[List[FileData], List[FileData]]:
    """
    Classify files into images and documents based on their MIME types.

    Args:
        files (List[FileData]): List of FileData objects to classify.

    Returns:
        Tuple[List[FileData], List[FileData]]: Two lists - one for image files, one for document files.
    """
    # Separate image files and document files by MIME type
    image_files = [f for f in files if f.type.startswith("image/")]
    doc_files = [f for f in files if f.type.startswith(("text/", "application/"))]
    return image_files, doc_files

async def handle_file_processing(content: str, files: List[FileData]) -> ProcessedMessage:
    """
    Process provided files, extract text, and return combined content with dspy.Image objects.

    Args:
        content (str): Additional content provided by the user.
        files (List[FileData]): List of FileData objects (already base64 or bytes).

    Returns:
        ProcessedMessage: Combined text content and list of dspy.Image objects.
    """
    if isinstance(content, str) and content.strip() == "":
        content = None
    
    if not files:
        return ProcessedMessage(
            content=content,
            images=None,
            context=None,
            recent_conversations=None
        )
    extracted_texts = []
    dspy_images = []

    # Classify files into images and documents
    image_files, doc_files = classify_file(files)

    # Extract text from document files
    for file_data in doc_files:
        text = extract_text(file_data)
        if text:
            extracted_texts.append(text)

    # Convert image files to dspy.Image objects
    for file_data in image_files:
        try:
            if isinstance(file_data.file, (bytes, bytearray)):
                base64_data = base64.b64encode(file_data.file).decode("utf-8")
            else:
                print(f"Unsupported file type for image {file_data.name}")
                continue

            dspy_image = await convert_to_dspy_image(base64_data)
            dspy_images.append(dspy_image)
        except Exception as e:
            print(f"Failed to convert image {file_data.name}: {e}")

    # Combine original content with extracted text
    combined_content = "\n\n".join(filter(None, [content] + extracted_texts))
    return ProcessedMessage(
        content=combined_content,
        images=dspy_images,
        context=None,
        recent_conversations=None
    )
